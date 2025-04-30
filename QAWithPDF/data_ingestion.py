import sys
from QAWithPDF.exception import customexception
from logger import logging
from llama_index.core import Document

import docx
import fitz  # PyMuPDF

def load_data(uploaded_file):
    """
    Load and parse text from uploaded .txt, .pdf, or .docx file.

    Parameters:
    - uploaded_file: Streamlit uploaded file (BytesIO)

    Returns:
    - List[Document]: A list containing one LlamaIndex Document object
    """
    try:
        logging.info(f"Loading file: {uploaded_file.name}")

        file_type = uploaded_file.name.split('.')[-1].lower()

        if file_type == "txt":
            content = uploaded_file.read().decode("utf-8")

        elif file_type == "pdf":
            pdf_doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            content = ""
            for page in pdf_doc:
                content += page.get_text()
            pdf_doc.close()

        elif file_type == "docx":
            doc = docx.Document(uploaded_file)
            content = "\n".join([para.text for para in doc.paragraphs])

        else:
            raise ValueError("Unsupported file type. Please upload a .txt, .pdf, or .docx file.")

        document = Document(text=content, metadata={"filename": uploaded_file.name})
        logging.info(f"File loaded successfully: {uploaded_file.name}")
        return [document]

    except Exception as e:
        logging.error("Error during document loading.")
        raise customexception(e, sys)
